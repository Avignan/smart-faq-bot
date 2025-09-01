from docx import Document
import os
import io
from pymongo import MongoClient
from dotenv import load_dotenv

load_dotenv()
mongodb_URI = os.getenv("MONGODB_URI")


def extract_text_from_docx(doc):
    try:
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from {doc}: {e}")
        return ""
    
def remove_headers_and_footers(doc):
    try:
        for section in doc.sections:
            for para in section.header.paragraphs:
                para.clear()
            for para in section.footer.paragraphs:
                para.clear()
        
        return doc
    except Exception as e:
        print(f"Error removing headers and footers from {doc}: {e}")
        return None


def preprocess_docx(file_bytes):
    try:
        file_stream = io.BytesIO(file_bytes)
        doc = Document(file_stream)
        doc_text = remove_headers_and_footers(doc)
        # doc = extract_text_from_docx(doc)
        return doc_text
    except Exception as e:
        print(f"Error pre-processing: {e}")
        return None
        

def db_connection(mongodb_URI):
    try:
        client = MongoClient(mongodb_URI)
        db = client["FAQ_RAG"]
        db = db['Document_Embedding_Info']
    except Exception as e:
        print(f"Error connecting to database: {e}")
        db = None
    return db